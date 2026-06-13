import json
import base64
from typing import Dict, Any, List
from pathlib import Path
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ReportGenerator:
    def generate(self, result: Dict[str, Any], format: str = "json") -> Any:
        if format == "json":
            return self.generate_json(result)
        elif format == "markdown":
            return self.generate_markdown(result)
        elif format == "word":
            return self.generate_word(result)
        else:
            raise ValueError(f"不支持的格式: {format}")
    
    def generate_json(self, result: Dict[str, Any]) -> str:
        return json.dumps(result, ensure_ascii=False, indent=2)
    
    def generate_markdown(self, result: Dict[str, Any]) -> str:
        lines = []
        lines.append("# 标书对比分析报告\n")
        lines.append(f"**生成时间:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        lines.append(f"**任务ID:** {result['task_id']}\n")

        lines.append("## 分析概览\n")
        lines.append(f"- **评审标准数量:** {result['requirements_count']}")
        lines.append(f"- **问题数量:** {result['issues_count']}")
        lines.append(f"- **综合得分:** {result['score']}分\n")

        # 第一块：评审标准
        lines.append("## 一、评审标准\n")
        for idx, c in enumerate(result['criteria'], 1):
            lines.append(f"### {idx}. {c['category']}\n")
            lines.append(f"**要求:** {c['requirement']}")
            lines.append(f"**位置:** {c['location']}")
            lines.append(f"**是否强制:** {'是' if c['is_mandatory'] else '否'}\n")
            lines.append("---\n")

        # 第二块：标书相关内容
        lines.append("## 二、标书相关内容\n")
        for idx, b in enumerate(result['bid_contents'], 1):
            lines.append(f"### 对应评审标准 {idx}\n")
            lines.append(f"**标书匹配内容:**\n")
            lines.append(f"> {b['matched_text']}\n")
            lines.append("---\n")

        # 第三块：对比结果
        lines.append("## 三、对比结果与建议\n")
        for idx, c in enumerate(result['comparisons'], 1):
            status = c.get('status', '未知')
            severity = c.get('severity', '未知')

            if status == '符合':
                icon = '✅'
            elif status == '部分符合':
                icon = '⚠️'
            else:
                icon = '❌'

            lines.append(f"### {icon} 评审标准 {idx} - {status}\n")
            lines.append(f"**严重程度:** {severity}\n")

            issues = c.get('issues', [])
            if issues:
                lines.append("**问题:**")
                for issue in issues:
                    lines.append(f"- {issue}")
                lines.append("")

            suggestions = c.get('suggestions', [])
            if suggestions:
                lines.append("**修改建议:**")
                for suggestion in suggestions:
                    lines.append(f"- {suggestion}")
                lines.append("")

            lines.append("---\n")

        return "\n".join(lines)
    
    def generate_word(self, result: Dict[str, Any]) -> bytes:
        doc = Document()

        title = doc.add_heading('标书对比分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph(f"任务ID: {result['task_id']}")

        doc.add_heading('分析概览', level=1)
        doc.add_paragraph(f"评审标准数量: {result['requirements_count']}")
        doc.add_paragraph(f"问题数量: {result['issues_count']}")
        doc.add_paragraph(f"综合得分: {result['score']}分")

        # 第一块：评审标准
        doc.add_heading('一、评审标准', level=1)
        for idx, c in enumerate(result['criteria'], 1):
            doc.add_heading(f"{idx}. {c['category']}", level=2)
            doc.add_paragraph(f"要求: {c['requirement']}")
            doc.add_paragraph(f"位置: {c['location']}")
            doc.add_paragraph(f"是否强制: {'是' if c['is_mandatory'] else '否'}")

        # 第二块：标书相关内容
        doc.add_heading('二、标书相关内容', level=1)
        for idx, b in enumerate(result['bid_contents'], 1):
            doc.add_heading(f"对应评审标准 {idx}", level=2)
            doc.add_paragraph(b['matched_text'])

        # 第三块：对比结果
        doc.add_heading('三、对比结果与建议', level=1)
        for idx, c in enumerate(result['comparisons'], 1):
            status = c.get('status', '未知')
            severity = c.get('severity', '未知')
            doc.add_heading(f"评审标准 {idx} - {status}", level=2)
            doc.add_paragraph(f"严重程度: {severity}")

            issues = c.get('issues', [])
            if issues:
                doc.add_paragraph("问题:", style='List Bullet')
                for issue in issues:
                    doc.add_paragraph(issue, style='List Bullet 2')

            suggestions = c.get('suggestions', [])
            if suggestions:
                doc.add_paragraph("修改建议:", style='List Bullet')
                for suggestion in suggestions:
                    doc.add_paragraph(suggestion, style='List Bullet 2')

            doc.add_paragraph()

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    def save_report(self, result: Dict[str, Any], format: str, output_dir: str = "tasks") -> str:
        Path(output_dir).mkdir(parents=True, exist_ok=True)
        
        task_id = result['task_id']
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if format == "json":
            filename = f"{task_id}_{timestamp}.json"
            filepath = Path(output_dir) / filename
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(self.generate_json(result))
        
        elif format == "markdown":
            filename = f"{task_id}_{timestamp}.md"
            filepath = Path(output_dir) / filename
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(self.generate_markdown(result))
        
        elif format == "word":
            filename = f"{task_id}_{timestamp}.docx"
            filepath = Path(output_dir) / filename
            with open(filepath, 'wb') as f:
                f.write(self.generate_word(result))
        
        else:
            raise ValueError(f"不支持的格式: {format}")
        
        return str(filepath)
